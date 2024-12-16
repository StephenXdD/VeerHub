import os
from docx import Document

def clear_document_content(doc_path, output_path):
    # Open the document
    doc = Document(doc_path)

    # Remove all paragraphs
    for _ in range(len(doc.paragraphs)):
        p = doc.paragraphs[0]
        p.clear()
        p._element.getparent().remove(p._element)
        
    # Remove all tables
    for _ in range(len(doc.tables)):
        tbl = doc.tables[0]
        tbl._element.getparent().remove(tbl._element)
    
    # Clear headers and footers
    for section in doc.sections:
        for header in section.header.paragraphs:
            header.clear()
            header._element.getparent().remove(header._element)
        for footer in section.footer.paragraphs:
            footer.clear()
            footer._element.getparent().remove(footer._element)

    # Ensure the output directory exists
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    
    # Save the cleared document
    doc.save(output_path)
    print(f"Cleared and saved document: {output_path}")

def process_folder(input_folder, output_folder):
    for root, _, files in os.walk(input_folder):
        for filename in files:
            if filename.lower().endswith('.docx') and not filename.startswith('~$'):
                input_file_path = os.path.join(root, filename)
                
                # Construct the output file path while preserving directory structure
                relative_path = os.path.relpath(root, input_folder)
                output_file_path = os.path.join(output_folder, relative_path, filename)

                # Clear the document content and save to the output folder
                clear_document_content(input_file_path, output_file_path)

# Paths to the input and output folders
input_folder = os.path.join(os.path.expanduser("~"), "Desktop", "trial_output")
output_folder = os.path.join(os.path.expanduser("~"), "Desktop", "modified_output")

# Process all .docx files in the folder
process_folder(input_folder, output_folder)
